import json
import re
from pathlib import Path
from typing import Dict, List

from docx import Document as DocxDocument
from lxml import etree
from pypdf import PdfReader
try:
    import pdfplumber
except Exception:
    pdfplumber = None

from langchain_core.documents import Document


class DocumentParser:
    @staticmethod
    def semantic_wrapper(tag_path: str, value: str, attributes: Dict[str, str] | None = None) -> str:
        clean_tag = tag_path.replace("/", " > ").strip()
        parts = [f"The field {clean_tag} is {value}."]
        if attributes:
            for key, attr_val in attributes.items():
                parts.append(f"For {clean_tag}, the attribute {key} is {attr_val}.")
        return " ".join(parts)

    @staticmethod
    def norm_key(key: str) -> str:
        return re.sub(r"[^a-z0-9]+", "_", key.strip().lower()).strip("_")

    @staticmethod
    def split_table_line(line: str) -> List[str]:
        if "|" in line:
            return [part.strip() for part in line.split("|") if part.strip()]
        if "\t" in line:
            return [part.strip() for part in line.split("\t") if part.strip()]
        if "," in line:
            return [part.strip() for part in line.split(",") if part.strip()]
        return [part.strip() for part in re.split(r"\s{2,}", line) if part.strip()]

    @staticmethod
    def split_loose_line(line: str) -> List[str]:
        normalized = re.sub(r"(?<=[A-Za-z])(?=\d)", " ", line.strip())
        normalized = re.sub(r"(?<=\d)(?=[A-Za-z])", " ", normalized)
        raw_tokens = [part.strip() for part in re.split(r"\s+", normalized) if part.strip()]
        merged_tokens: List[str] = []
        i = 0
        while i < len(raw_tokens):
            current = raw_tokens[i]
            nxt = raw_tokens[i + 1] if i + 1 < len(raw_tokens) else ""
            if re.fullmatch(r"\d{4}-\d{2}-\d{2}", current) and re.fullmatch(r"\d{2}:\d{2}(?::\d{2})?", nxt):
                merged_tokens.append(f"{current} {nxt}")
                i += 2
                continue
            merged_tokens.append(current)
            i += 1
        return merged_tokens

    def infer_header_columns(self, text: str) -> List[str]:
        lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
        for line in lines[:12]:
            parts = self.split_table_line(line)
            if self.looks_like_generic_header(parts):
                return [p.strip() for p in parts]
        return []

    def table_rows_to_records(self, rows: List[List[str]]) -> List[Dict[str, str]]:
        if not rows:
            return []
        cleaned: List[List[str]] = []
        for row in rows:
            cells = [("" if c is None else str(c).strip()) for c in row]
            if any(cells):
                cleaned.append(cells)
        if not cleaned:
            return []

        header_idx = -1
        headers: List[str] = []
        for idx, row in enumerate(cleaned[:3]):
            alpha = sum(1 for cell in row if re.search(r"[A-Za-z]", cell))
            numeric = sum(1 for cell in row if re.fullmatch(r"-?\d+(?:\.\d+)?", cell))
            if alpha >= max(2, len(row) // 2) and numeric <= max(1, len(row) // 5):
                headers = [cell if cell else f"col_{i + 1}" for i, cell in enumerate(row)]
                header_idx = idx
                break
        if not headers:
            max_cols = max(len(r) for r in cleaned)
            headers = [f"col_{i + 1}" for i in range(max_cols)]
            header_idx = -1

        records: List[Dict[str, str]] = []
        start = header_idx + 1 if header_idx >= 0 else 0
        for row in cleaned[start:]:
            if len(row) < len(headers):
                row = row + [""] * (len(headers) - len(row))
            if len(row) > len(headers):
                row = row[: len(headers)]
            rec = {headers[i].strip(): row[i].strip() for i in range(len(headers))}
            rec["raw_record"] = " ".join([v for v in row if v])
            if any(v for v in rec.values()):
                records.append(rec)
        return records

    def extract_pdfplumber_records(self, file_path: str) -> Dict[int, List[Dict[str, str]]]:
        per_page: Dict[int, List[Dict[str, str]]] = {}
        if pdfplumber is None:
            return per_page
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_records: List[Dict[str, str]] = []
                    tables = page.extract_tables() or []
                    for table in tables:
                        rows = table or []
                        page_records.extend(self.table_rows_to_records(rows))
                    if page_records:
                        per_page[page_num] = page_records
        except Exception:
            return {}
        return per_page

    @staticmethod
    def looks_like_generic_header(parts: List[str]) -> bool:
        if len(parts) < 2:
            return False
        alpha_tokens = 0
        numeric_tokens = 0
        for p in parts:
            if re.search(r"[a-zA-Z]", p):
                alpha_tokens += 1
            if re.fullmatch(r"-?\d+(?:\.\d+)?", p.strip()):
                numeric_tokens += 1
        return alpha_tokens >= max(2, len(parts) // 2) and numeric_tokens <= max(1, len(parts) // 5)

    def extract_table_records(self, text: str) -> List[Dict[str, str]]:
        lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
        if not lines:
            return []
        header_idx = -1
        headers: List[str] = []
        for idx, line in enumerate(lines[:12]):
            parts = self.split_table_line(line)
            if self.looks_like_generic_header(parts):
                headers = parts
                header_idx = idx
                break
        if header_idx == -1 or not headers:
            return []

        records: List[Dict[str, str]] = []
        for line in lines[header_idx + 1 :]:
            if re.fullmatch(r"[-=_\s|]+", line):
                continue
            cols = self.split_table_line(line)
            if len(cols) < 2:
                continue
            if len(cols) < len(headers):
                cols.extend([""] * (len(headers) - len(cols)))
            if len(cols) > len(headers):
                cols = cols[: len(headers) - 1] + [" ".join(cols[len(headers) - 1 :])]
            rec = {headers[i].strip(): cols[i].strip() for i in range(len(headers))}
            if any(v for v in rec.values()):
                records.append(rec)
        return records

    def extract_line_records(self, text: str, expected_cols: int = 0, headers: List[str] | None = None) -> List[Dict[str, str]]:
        """Fallback parser for row-like text when no explicit table header exists."""
        lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
        if not lines:
            return []

        def chunk_tokens(tokens: List[str]) -> List[List[str]]:
            if not tokens:
                return []
            chunks: List[List[str]] = []
            # If we have an expected row width from header, use it first.
            if expected_cols and len(tokens) >= expected_cols * 2:
                i = 0
                while i < len(tokens):
                    chunk = tokens[i : i + expected_cols]
                    if len(chunk) >= max(4, expected_cols // 2):
                        chunks.append(chunk)
                    i += expected_cols
                return chunks

            # Generic fallback: split long lines into records on numeric-tail boundaries.
            current: List[str] = []
            numeric_seen = 0
            for idx, token in enumerate(tokens):
                current.append(token)
                if re.fullmatch(r"-?\d+(?:\.\d+)?", token):
                    numeric_seen += 1
                next_tok = tokens[idx + 1] if idx + 1 < len(tokens) else ""
                if (
                    len(current) >= 6
                    and numeric_seen >= 2
                    and next_tok
                    and re.search(r"[A-Za-z]", next_tok)
                    and re.fullmatch(r"-?\d+(?:\.\d+)?", token)
                ):
                    chunks.append(current)
                    current = []
                    numeric_seen = 0
            if current:
                chunks.append(current)
            return chunks

        records: List[Dict[str, str]] = []
        for line in lines:
            if re.fullmatch(r"[-=_\s|]+", line):
                continue
            tokens = self.split_loose_line(line)
            for chunk in chunk_tokens(tokens):
                if len(chunk) < 4:
                    continue
                has_alpha = any(re.search(r"[A-Za-z]", t) for t in chunk)
                has_digit = any(re.search(r"\d", t) for t in chunk)
                if not (has_alpha and has_digit):
                    continue
                rec: Dict[str, str] = {"raw_record": " ".join(chunk)}
                for idx, token in enumerate(chunk[:14], start=1):
                    rec[f"col_{idx}"] = token
                if headers:
                    for idx, token in enumerate(chunk[: len(headers)]):
                        rec[self.norm_key(headers[idx])] = token
                numbers = [t for t in chunk if re.fullmatch(r"-?\d+(?:\.\d+)?", t)]
                for idx, value in enumerate(numbers[:8], start=1):
                    rec[f"numeric_{idx}"] = value
                records.append(rec)
        return records

    def record_to_doc(self, record: Dict[str, str], source_name: str, page_num: int, idx: int) -> Document:
        normalized = {self.norm_key(k): v for k, v in record.items()}

        page_content = " ".join(f"The {k} is {v}." for k, v in record.items() if v)
        metadata = {
            "source": source_name,
            "type": "pdf",
            "page": page_num,
            "doc_kind": "categorical_row",
            "row_id": f"{page_num}-{idx}",
            "row_payload": " | ".join(f"{k}: {v}" for k, v in record.items()),
            "row_fields_json": json.dumps(normalized, ensure_ascii=True),
        }
        return Document(page_content=page_content, metadata=metadata)

    def load_pdf(self, file_path: str, source_name: str) -> List[Document]:
        reader = PdfReader(file_path)
        table_records_by_page = self.extract_pdfplumber_records(file_path)
        docs: List[Document] = []
        header_context: List[str] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                continue
            docs.append(Document(page_content=text, metadata={"source": source_name, "type": "pdf", "page": page_num}))

            records = table_records_by_page.get(page_num, [])
            if not records:
                headers = self.infer_header_columns(text)
                if headers:
                    header_context = headers
                elif header_context:
                    headers = header_context
                table_records = self.extract_table_records(text)
                line_records = self.extract_line_records(text, expected_cols=len(headers), headers=headers)
                # Merge both parse strategies to maximize recall, then de-duplicate.
                merged = table_records + line_records
                records = []
                seen = set()
                for rec in merged:
                    normalized = {self.norm_key(k): str(v).strip() for k, v in rec.items() if str(v).strip()}
                    # Prefer stable raw_record key when present; otherwise all normalized fields.
                    row_key = normalized.get("raw_record", "") or json.dumps(
                        normalized,
                        sort_keys=True,
                        ensure_ascii=True,
                    )
                    if row_key in seen:
                        continue
                    seen.add(row_key)
                    records.append(rec)

            for idx, rec in enumerate(records, start=1):
                docs.append(self.record_to_doc(rec, source_name, page_num, idx))
        return docs

    def load_docx(self, file_path: str, source_name: str) -> List[Document]:
        doc = DocxDocument(file_path)
        text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text and p.text.strip())
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": source_name, "type": "docx"})]

    def load_xml(self, file_path: str, source_name: str) -> List[Document]:
        parser = etree.XMLParser(recover=True, remove_blank_text=True)
        tree = etree.parse(file_path, parser=parser)
        root = tree.getroot()
        lines: List[str] = []

        def walk(node: etree._Element, path: str) -> None:
            current_path = f"{path}/{node.tag}" if path else str(node.tag)
            node_text = (node.text or "").strip()
            if node_text:
                lines.append(self.semantic_wrapper(current_path, node_text, dict(node.attrib) or None))
            elif node.attrib:
                for k, v in node.attrib.items():
                    lines.append(f"The field {current_path} has attribute {k} with value {v}.")
            for child in node:
                walk(child, current_path)

        walk(root, "")
        if not lines:
            return []
        return [Document(page_content="\n".join(lines), metadata={"source": source_name, "type": "xml"})]

    def load_documents(self, file_path: str, source_name: str) -> List[Document]:
        ext = Path(source_name).suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(file_path, source_name)
        if ext == ".docx":
            return self.load_docx(file_path, source_name)
        if ext == ".xml":
            return self.load_xml(file_path, source_name)
        raise ValueError(f"Unsupported file type: {ext}")
