import json
import re
from pathlib import Path
from typing import Dict, List

from docx import Document as DocxDocument
from lxml import etree
from pypdf import PdfReader

from langchain_core.documents import Document


class DocumentParser:
    @staticmethod
    def semantic_wrapper(tag_path: str, value: str, attributes: Dict[str, str] | None = None) -> str:
        clean_tag = tag_path.replace("/", " > ").strip()
        parts = [f"The field {clean_tag} is {value}."]
        if attributes:
            for key, attr_val in attributes.items():
                parts.append(f"For {clean_tag}, the attribute {key} is {attr_val}.")
        return " ".join(parts)

    @staticmethod
    def norm_key(key: str) -> str:
        return re.sub(r"[^a-z0-9]+", "_", key.strip().lower()).strip("_")

    @staticmethod
    def split_table_line(line: str) -> List[str]:
        if "|" in line:
            return [part.strip() for part in line.split("|") if part.strip()]
        if "\t" in line:
            return [part.strip() for part in line.split("\t") if part.strip()]
        if "," in line:
            return [part.strip() for part in line.split(",") if part.strip()]
        return [part.strip() for part in re.split(r"\s{2,}", line) if part.strip()]

    def looks_like_header(self, line: str) -> bool:
        low = line.lower()
        has_machine = "machine" in low or re.search(r"\bm_id\b|\bmachine_id\b", low) is not None
        has_job = "job" in low or re.search(r"\bjob_id\b", low) is not None
        return bool(has_machine and has_job)

    def extract_table_records(self, text: str) -> List[Dict[str, str]]:
        lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
        if not lines:
            return []
        header_idx = -1
        headers: List[str] = []
        for idx, line in enumerate(lines):
            if self.looks_like_header(line):
                parts = self.split_table_line(line)
                if len(parts) >= 2:
                    headers = parts
                    header_idx = idx
                    break
        if header_idx == -1 or not headers:
            return []

        records: List[Dict[str, str]] = []
        for line in lines[header_idx + 1 :]:
            if re.fullmatch(r"[-=_\s|]+", line):
                continue
            cols = self.split_table_line(line)
            if len(cols) < 2:
                continue
            if len(cols) < len(headers):
                cols.extend([""] * (len(headers) - len(cols)))
            if len(cols) > len(headers):
                cols = cols[: len(headers) - 1] + [" ".join(cols[len(headers) - 1 :])]
            rec = {headers[i].strip(): cols[i].strip() for i in range(len(headers))}
            if any(v for v in rec.values()):
                records.append(rec)
        return records

    def extract_machine_line_records(self, text: str) -> List[Dict[str, str]]:
        records: List[Dict[str, str]] = []
        for line in text.splitlines():
            clean = line.strip()
            if not clean:
                continue
            machine_match = re.search(r"\bM\d{1,4}\b", clean, flags=re.IGNORECASE)
            if not machine_match:
                continue
            job_match = re.search(r"\b(JOB[_\-\s]?\d+|J\d{1,6}|JOB[A-Z0-9_\-]+)\b", clean, flags=re.IGNORECASE)
            rec: Dict[str, str] = {"Machine": machine_match.group(0).upper(), "RawLine": clean}
            if job_match:
                rec["Job"] = job_match.group(0).upper()
            records.append(rec)
        return records

    @staticmethod
    def normalize_datetime_token(token: str) -> str:
        s = re.sub(r"\s*-\s*", "-", token)
        s = re.sub(r"\s*:\s*", ":", s)
        return re.sub(r"\s+", " ", s).strip()

    def extract_job_segment_records(self, text: str) -> List[Dict[str, str]]:
        compact = re.sub(r"\s+", " ", text or "").strip()
        starts = list(re.finditer(r"\bJ\d{2,6}\b", compact, flags=re.IGNORECASE))
        if not starts:
            return []

        out: List[Dict[str, str]] = []
        for i, match in enumerate(starts):
            st = match.start()
            en = starts[i + 1].start() if i + 1 < len(starts) else len(compact)
            seg = compact[st:en].strip()
            job_m = re.search(r"\bJ\d{2,6}\b", seg, flags=re.IGNORECASE)
            machine_m = re.search(r"\bM\d{1,4}\b", seg, flags=re.IGNORECASE)
            if not (job_m and machine_m):
                continue
            rec: Dict[str, str] = {
                "Job": job_m.group(0).upper(),
                "Machine": machine_m.group(0).upper(),
                "RawLine": seg,
            }

            seg_after_machine = seg[machine_m.end() :].strip()
            first_dt = re.search(
                r"\d{4}\s*-\s*\d{2}\s*-\s*\d{2}\s+\d{2}\s*:\s*\d{2}\s*:\s*\d{2}",
                seg_after_machine,
                flags=re.IGNORECASE,
            )
            prefix = seg_after_machine[: first_dt.start()].strip() if first_dt else seg_after_machine
            alpha = re.sub(r"[^a-z]", "", prefix.lower())
            for op, aliases in [
                ("Grinding", ["grinding", "grind"]),
                ("Milling", ["milling", "milli"]),
                ("Lathe", ["lathe"]),
                ("Drilling", ["drilling", "drill"]),
                ("Additive", ["additive"]),
            ]:
                if any(a in alpha for a in aliases):
                    rec["Operation_Type"] = op
                    break
            nums = re.findall(r"\d+(?:\.\d+)?", prefix)
            if len(nums) >= 1:
                rec["Material_Used"] = nums[0]
            if len(nums) >= 2:
                rec["Processing_Time"] = nums[1]
            if len(nums) >= 3:
                rec["Energy_Consumption"] = nums[2]
            if len(nums) >= 4:
                rec["Machine_Availability"] = nums[3]

            dt_tokens = re.findall(
                r"\d{4}\s*-\s*\d{2}\s*-\s*\d{2}\s+\d{2}\s*:\s*\d{2}\s*:\s*\d{2}",
                seg,
                flags=re.IGNORECASE,
            )
            dt_tokens = [self.normalize_datetime_token(t) for t in dt_tokens]
            if len(dt_tokens) >= 1:
                rec["Start_Time"] = dt_tokens[0]
            if len(dt_tokens) >= 2:
                rec["End_Time"] = dt_tokens[1]
            if len(dt_tokens) >= 3:
                rec["Actual_Start"] = dt_tokens[2]
            if len(dt_tokens) >= 4:
                rec["Actual_End"] = dt_tokens[3]

            if re.search(r"c\s*o\s*m\s*p\s*l\s*e\s*t\s*e\s*d", seg, flags=re.IGNORECASE):
                rec["Job_Status"] = "Completed"
            elif re.search(r"f\s*a\s*i\s*l\s*e\s*d", seg, flags=re.IGNORECASE):
                rec["Job_Status"] = "Failed"
            elif re.search(r"d\s*e\s*l\s*a\s*y\s*e?\s*d", seg, flags=re.IGNORECASE):
                rec["Job_Status"] = "Delayed"
            out.append(rec)
        return out

    def record_to_doc(self, record: Dict[str, str], source_name: str, page_num: int, idx: int) -> Document:
        normalized = {self.norm_key(k): v for k, v in record.items()}
        machine = ""
        job_type = ""
        job_value = ""
        for k, v in normalized.items():
            cand = str(v).strip()
            if k in {"machine", "machine_id", "m_id"} and re.fullmatch(r"M\d{1,4}", cand.upper()):
                machine = cand.upper()
            if ("job" in k and "type" in k) or k in {"job_type", "type_of_job", "type", "category", "process_type", "operation_type"}:
                job_type = cand
            if "job" in k and "type" not in k and k not in {"job_count"} and re.fullmatch(r"J\d{2,6}", cand.upper()):
                job_value = cand.upper()

        page_content = " ".join(f"The {k} is {v}." for k, v in record.items() if v)
        metadata = {
            "source": source_name,
            "type": "pdf",
            "page": page_num,
            "doc_kind": "categorical_row",
            "row_id": f"{page_num}-{idx}",
            "row_machine": machine,
            "row_job_type": job_type,
            "row_job": job_value,
            "row_payload": " | ".join(f"{k}: {v}" for k, v in record.items()),
            "row_fields_json": json.dumps(normalized, ensure_ascii=True),
        }
        return Document(page_content=page_content, metadata=metadata)

    def load_pdf(self, file_path: str, source_name: str) -> List[Document]:
        reader = PdfReader(file_path)
        docs: List[Document] = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if not text.strip():
                continue
            docs.append(Document(page_content=text, metadata={"source": source_name, "type": "pdf", "page": page_num}))

            records = self.extract_table_records(text)
            if not records:
                records = self.extract_machine_line_records(text)
            records.extend(self.extract_job_segment_records(text))

            deduped: List[Dict[str, str]] = []
            seen = set()
            for rec in records:
                key = (
                    str(rec.get("Job", "")).upper(),
                    str(rec.get("Machine", "")).upper(),
                    str(rec.get("Start_Time", "")),
                    str(rec.get("End_Time", "")),
                    str(rec.get("RawLine", ""))[:160],
                )
                if key in seen:
                    continue
                seen.add(key)
                deduped.append(rec)
            for idx, rec in enumerate(deduped, start=1):
                docs.append(self.record_to_doc(rec, source_name, page_num, idx))
        return docs

    def load_docx(self, file_path: str, source_name: str) -> List[Document]:
        doc = DocxDocument(file_path)
        text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text and p.text.strip())
        if not text:
            return []
        return [Document(page_content=text, metadata={"source": source_name, "type": "docx"})]

    def load_xml(self, file_path: str, source_name: str) -> List[Document]:
        parser = etree.XMLParser(recover=True, remove_blank_text=True)
        tree = etree.parse(file_path, parser=parser)
        root = tree.getroot()
        lines: List[str] = []

        def walk(node: etree._Element, path: str) -> None:
            current_path = f"{path}/{node.tag}" if path else str(node.tag)
            node_text = (node.text or "").strip()
            if node_text:
                lines.append(self.semantic_wrapper(current_path, node_text, dict(node.attrib) or None))
            elif node.attrib:
                for k, v in node.attrib.items():
                    lines.append(f"The field {current_path} has attribute {k} with value {v}.")
            for child in node:
                walk(child, current_path)

        walk(root, "")
        if not lines:
            return []
        return [Document(page_content="\n".join(lines), metadata={"source": source_name, "type": "xml"})]

    def load_documents(self, file_path: str, source_name: str) -> List[Document]:
        ext = Path(source_name).suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(file_path, source_name)
        if ext == ".docx":
            return self.load_docx(file_path, source_name)
        if ext == ".xml":
            return self.load_xml(file_path, source_name)
        raise ValueError(f"Unsupported file type: {ext}")
